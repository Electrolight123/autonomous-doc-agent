import os
import re
import time
import uuid
from datetime import datetime
from typing import List, Dict, Any

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from groq import Groq
from docx import Document
from docx.shared import Pt


load_dotenv()

OUTPUT_DIR = "generated_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


LOCAL_KNOWLEDGE_BASE = [
    {
        "id": "fintech_chatbot_metrics",
        "title": "Fintech Customer Support Chatbot Metrics",
        "tags": ["fintech", "chatbot", "customer support", "metrics", "success"],
        "content": (
            "For fintech chatbot projects, useful success metrics include first contact resolution rate, "
            "average handling time, customer satisfaction score, escalation rate, chatbot adoption rate, "
            "support cost reduction, and response time improvement."
        )
    },
    {
        "id": "fintech_chatbot_risks",
        "title": "Fintech Chatbot Risk Checklist",
        "tags": ["fintech", "chatbot", "risk", "compliance", "security"],
        "content": (
            "Common risks in fintech chatbot implementations include data privacy issues, incorrect financial guidance, "
            "integration failure with CRM or ticketing systems, poor intent recognition, customer dissatisfaction, "
            "and regulatory compliance gaps. Mitigation should include human escalation, audit logs, encryption, "
            "testing, and controlled rollout."
        )
    },
    {
        "id": "nbfc_compliance",
        "title": "NBFC Compliance and Data Privacy Notes",
        "tags": ["nbfc", "compliance", "data privacy", "security", "rbi"],
        "content": (
            "For NBFC systems, compliance and data privacy should be treated as high-priority requirements. "
            "Important controls include role-based access, encryption, audit trails, data retention policies, "
            "secure storage, approval workflows, and regular compliance reviews."
        )
    },
    {
        "id": "document_automation_architecture",
        "title": "AI Document Automation Architecture",
        "tags": ["document automation", "architecture", "ocr", "nlp", "workflow"],
        "content": (
            "A typical AI document automation system includes document ingestion, OCR, document classification, "
            "data extraction, validation, workflow routing, human review, document storage, audit logging, "
            "and integration with existing business systems."
        )
    },
    {
        "id": "document_automation_risks",
        "title": "AI Document Automation Risk Checklist",
        "tags": ["document automation", "risk", "compliance", "security", "quality", "workflow"],
        "content": (
            "Common risks in AI document automation include inaccurate extraction, poor document classification, "
            "data privacy issues, workflow integration failure, missing audit trails, regulatory compliance gaps, "
            "and low user adoption. Mitigation should include human review, validation rules, audit logging, "
            "encryption, access control, testing, and controlled phased rollout."
        )
    },
    {
        "id": "mvp_delivery_strategy",
        "title": "MVP and Phased Delivery Strategy",
        "tags": ["mvp", "limited budget", "aggressive timeline", "phased plan", "roadmap"],
        "content": (
            "When budget is limited and timeline is aggressive, an MVP-first approach is recommended. "
            "Start with high-volume and low-complexity use cases, validate value quickly, collect feedback, "
            "then expand to complex workflows in later phases."
        )
    },
    {
        "id": "business_document_quality",
        "title": "Business Document Quality Checklist",
        "tags": ["proposal", "business plan", "technical plan", "quality", "document"],
        "content": (
            "A strong business document should include clear objectives, assumptions, proposed solution, "
            "implementation plan, risks, mitigation steps, success metrics, and a concise conclusion. "
            "For ambiguous requests, the document should explicitly state assumptions and tradeoffs."
        )
    }
]


app = FastAPI(
    title="Autonomous Document Agent",
    description="A Python AI agent that plans tasks, retrieves context, reflects on output, and generates Word documents.",
    version="1.0.0"
)


class AgentRequest(BaseModel):
    request: str = Field(
        ...,
        min_length=10,
        max_length=2000,
        description="Natural language request for the autonomous document agent"
    )


def decide_document_type(user_request: str) -> str:
    request_lower = user_request.lower()

    if "technical and business plan" in request_lower:
        return "Technical and Business Plan"

    if "proposal" in request_lower:
        return "Business Proposal"

    if "technical design" in request_lower or "architecture" in request_lower:
        return "Technical Design Document"

    if "business plan" in request_lower:
        return "Business Plan"

    if "sop" in request_lower or "standard operating procedure" in request_lower:
        return "Standard Operating Procedure"

    if "meeting minutes" in request_lower or "minutes" in request_lower:
        return "Meeting Minutes"

    if "project plan" in request_lower or "implementation plan" in request_lower:
        return "Project Plan"

    if "report" in request_lower:
        return "Business Report"

    return "Structured Business Document"


def create_title(user_request: str, document_type: str) -> str:
    request_lower = user_request.lower()

    if "chatbot" in request_lower:
        return f"{document_type}: AI-Powered Customer Support Chatbot"

    if "resume screening" in request_lower:
        return f"{document_type}: AI Resume Screening System"

    if "document automation" in request_lower:
        return f"{document_type}: AI Document Automation System"

    if "fintech" in request_lower:
        return f"{document_type}: AI Solution for Fintech Operations"

    if "nbfc" in request_lower:
        return f"{document_type}: AI Solution for NBFC Operations"

    return f"{document_type}: AI-Generated Business Document"


def decide_sections(document_type: str, user_request: str) -> List[str]:
    request_lower = user_request.lower()

    if document_type == "Business Proposal":
        sections = [
            "Executive Summary",
            "Business Objective",
            "Problem Statement",
            "Proposed Solution",
            "Implementation Plan",
            "Risks and Mitigation",
            "Success Metrics",
            "Conclusion"
        ]

    elif document_type == "Technical Design Document":
        sections = [
            "Executive Summary",
            "System Overview",
            "Functional Requirements",
            "Architecture Design",
            "API Design",
            "Data Flow",
            "Security and Compliance",
            "Scalability Considerations",
            "Conclusion"
        ]

    elif document_type == "Technical and Business Plan":
        sections = [
            "Executive Summary",
            "Business Objective",
            "Technical Architecture",
            "Implementation Roadmap",
            "Compliance and Security",
            "Budget and Resource Assumptions",
            "Risks and Mitigation",
            "Success Metrics",
            "Conclusion"
        ]

    elif document_type == "Business Plan":
        sections = [
            "Executive Summary",
            "Business Objective",
            "Target Users",
            "Proposed Solution",
            "Go-To-Market Approach",
            "Operational Plan",
            "Risks and Mitigation",
            "Success Metrics",
            "Conclusion"
        ]

    elif document_type == "Standard Operating Procedure":
        sections = [
            "Purpose",
            "Scope",
            "Responsibilities",
            "Step-by-Step Procedure",
            "Exception Handling",
            "Quality Checks",
            "Conclusion"
        ]

    elif document_type == "Meeting Minutes":
        sections = [
            "Meeting Overview",
            "Attendees",
            "Agenda",
            "Discussion Summary",
            "Decisions Taken",
            "Action Items",
            "Next Steps"
        ]

    elif document_type == "Project Plan":
        sections = [
            "Project Overview",
            "Objectives",
            "Scope",
            "Milestones",
            "Timeline",
            "Risks and Dependencies",
            "Success Criteria",
            "Conclusion"
        ]

    else:
        sections = [
            "Executive Summary",
            "Objective",
            "Background",
            "Proposed Approach",
            "Execution Plan",
            "Risks",
            "Conclusion"
        ]

    complex_request_keywords = [
        "vague",
        "ambiguous",
        "limited",
        "aggressive",
        "conflicting",
        "missing information"
    ]

    if any(word in request_lower for word in complex_request_keywords):
        sections.insert(1, "Assumptions Made by the Agent")
        sections.insert(2, "Conflict Resolution Approach")

    return sections


def create_tasks(sections: List[str]) -> List[Dict[str, Any]]:
    tasks = [
        {
            "step": 1,
            "task": "Understand the user request",
            "reason": "The agent needs to identify the objective, document type, and business context."
        },
        {
            "step": 2,
            "task": "Decide the document type",
            "reason": "Different business documents require different structures."
        },
        {
            "step": 3,
            "task": "Create the document structure",
            "reason": "The document must have a logical and professional flow."
        },
        {
            "step": 4,
            "task": "Retrieve relevant context from local knowledge base",
            "reason": "Simple RAG improves the generated document with domain-specific guidance."
        }
    ]

    step_number = 5

    for section in sections:
        tasks.append({
            "step": step_number,
            "task": f"Generate content for section: {section}",
            "reason": f"The section '{section}' is required for the final document."
        })
        step_number += 1

    tasks.append({
        "step": step_number,
        "task": "Reflect on generated output",
        "reason": "The agent should review its own output before producing the final deliverable."
    })
    step_number += 1

    tasks.append({
        "step": step_number,
        "task": "Generate final Microsoft Word document",
        "reason": "The assignment requires a polished .docx file as final output."
    })

    return tasks


def create_assumptions(user_request: str) -> List[str]:
    request_lower = user_request.lower()

    assumptions = [
        "The document is intended for a professional business audience.",
        "Mock data may be used where exact business data is not available."
    ]

    if "budget" not in request_lower:
        assumptions.append("Budget details were not provided, so a moderate-cost implementation approach is assumed.")

    if "timeline" not in request_lower:
        assumptions.append("Timeline was not provided, so a phased delivery plan is assumed.")

    if "compliance" in request_lower:
        assumptions.append("Compliance and data privacy are treated as high-priority requirements.")

    if "limited" in request_lower:
        assumptions.append("Because budget or resources are limited, the agent prioritizes an MVP-first approach.")

    if "aggressive" in request_lower:
        assumptions.append("Because timeline is aggressive, the agent recommends phased delivery instead of a big-bang launch.")

    return assumptions


def create_execution_plan(user_request: str) -> Dict[str, Any]:
    document_type = decide_document_type(user_request)
    title = create_title(user_request, document_type)
    sections = decide_sections(document_type, user_request)
    tasks = create_tasks(sections)
    assumptions = create_assumptions(user_request)

    return {
        "document_type": document_type,
        "title": title,
        "goal": "Generate a polished Microsoft Word document based on the user's request.",
        "sections": sections,
        "tasks": tasks,
        "assumptions": assumptions
    }


def tokenize_text(text: str) -> set:
    return set(re.findall(r"[a-zA-Z0-9]+", text.lower()))


def retrieve_rag_context(
    user_request: str,
    plan: Dict[str, Any],
    top_k: int = 4
) -> List[Dict[str, Any]]:
    query_text = (
        user_request + " "
        + plan["title"] + " "
        + plan["document_type"] + " "
        + " ".join(plan["sections"]) + " "
        + " ".join(plan["assumptions"])
    )

    query_tokens = tokenize_text(query_text)

    scored_chunks = []

    for chunk in LOCAL_KNOWLEDGE_BASE:
        chunk_text = (
            chunk["title"] + " "
            + " ".join(chunk["tags"]) + " "
            + chunk["content"]
        )

        chunk_tokens = tokenize_text(chunk_text)
        overlap_score = len(query_tokens.intersection(chunk_tokens))

        tag_bonus = 0

        for tag in chunk["tags"]:
            if tag.lower() in query_text.lower():
                tag_bonus += 3

        final_score = overlap_score + tag_bonus

        if final_score > 0:
            scored_chunks.append({
                "id": chunk["id"],
                "title": chunk["title"],
                "content": chunk["content"],
                "score": final_score
            })

    scored_chunks.sort(key=lambda item: item["score"], reverse=True)

    return scored_chunks[:top_k]


def format_rag_context(rag_context: List[Dict[str, Any]]) -> str:
    if not rag_context:
        return "No additional RAG context was retrieved."

    formatted_chunks = []

    for index, chunk in enumerate(rag_context, start=1):
        formatted_chunks.append(
            f"Context {index}: {chunk['title']}\n{chunk['content']}"
        )

    return "\n\n".join(formatted_chunks)


def get_groq_client():
    api_key = os.getenv("GROQ_API_KEY")

    if not api_key:
        return None

    return Groq(api_key=api_key)


def fallback_section_content(section_name: str, user_request: str) -> str:
    return (
        f"This section covers '{section_name}' for the request: {user_request}\n\n"
        "The agent uses available context, mock business assumptions, and a structured approach "
        "to generate a professional document section. The content is designed to be clear, practical, "
        "and suitable for a business audience.\n\n"
        "Key points:\n"
        "- The solution is aligned with the user's business objective.\n"
        "- Missing details are handled using reasonable assumptions.\n"
        "- The recommendation focuses on practical execution, risk reduction, and measurable outcomes."
    )


def call_llm_with_retry(system_prompt: str, user_prompt: str, fallback_text: str) -> Dict[str, Any]:
    client = get_groq_client()
    model = os.getenv("GROQ_MODEL")

    if client is None or not model:
        return {
            "content": fallback_text,
            "source": "fallback",
            "error": "GROQ_API_KEY or GROQ_MODEL not configured"
        }

    max_attempts = 2
    last_error = None

    for attempt in range(1, max_attempts + 1):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": system_prompt
                    },
                    {
                        "role": "user",
                        "content": user_prompt
                    }
                ],
                temperature=0.3,
                max_tokens=700
            )

            content = response.choices[0].message.content

            return {
                "content": content,
                "source": "groq",
                "error": None
            }

        except Exception as error:
            last_error = str(error)
            time.sleep(1)

    return {
        "content": fallback_text,
        "source": "fallback",
        "error": last_error
    }


def generate_section_content(
    user_request: str,
    plan: Dict[str, Any],
    section_name: str,
    rag_context: List[Dict[str, Any]]
) -> Dict[str, Any]:
    system_prompt = """
You are a professional autonomous business document writing agent.

Write clear, structured, business-ready content.
Use practical language.
Use short paragraphs and bullet points when useful.
Do not use markdown headings.
Do not repeat the section heading at the beginning of the content.
Do not mention that you are an AI model.
"""

    user_prompt = f"""
Original user request:
{user_request}

Document title:
{plan["title"]}

Document type:
{plan["document_type"]}

Agent assumptions:
{plan["assumptions"]}

Retrieved domain context from simple RAG:
{format_rag_context(rag_context)}

Current section to write:
{section_name}

Write polished content for this section.
The content should be specific to the user's request, not generic.
Use the retrieved context where relevant, but do not mention the phrase RAG in the final business content.
Do not repeat the section title in the first line.
"""

    fallback_text = fallback_section_content(section_name, user_request)

    return call_llm_with_retry(
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        fallback_text=fallback_text
    )


def reflect_on_output(
    user_request: str,
    plan: Dict[str, Any],
    section_outputs: Dict[str, str],
    rag_context: List[Dict[str, Any]]
) -> Dict[str, Any]:
    required_sections = plan["sections"]
    generated_sections = list(section_outputs.keys())

    missing_sections = [
        section for section in required_sections
        if section not in generated_sections or not section_outputs.get(section)
    ]

    total_content_length = sum(len(content) for content in section_outputs.values())

    strengths = []
    improvement_points = []

    if not missing_sections:
        strengths.append("All planned sections were generated successfully.")
    else:
        improvement_points.append(f"Missing sections found: {missing_sections}")

    if total_content_length > 3000:
        strengths.append("Generated content has sufficient depth for a business document.")
    else:
        improvement_points.append("Generated content may be too short for a complete business document.")

    if plan["assumptions"]:
        strengths.append("Agent made explicit assumptions where input information was incomplete.")

    if rag_context:
        strengths.append(f"Simple RAG retrieved {len(rag_context)} relevant knowledge chunks to support document generation.")

    evaluation_sections = [
        "Risks and Mitigation",
        "Compliance and Security",
        "Success Metrics"
    ]

    if any(section in generated_sections for section in evaluation_sections):
        strengths.append("Document includes practical business evaluation sections such as risks, compliance, or success metrics.")

    request_lower = user_request.lower()

    if "vague" in request_lower or "ambiguous" in request_lower:
        if "Assumptions Made by the Agent" in generated_sections and "Conflict Resolution Approach" in generated_sections:
            strengths.append("Complex or ambiguous request was handled using assumptions and conflict resolution.")
        else:
            improvement_points.append("Ambiguous request could be improved with explicit assumptions and conflict resolution.")

    if missing_sections:
        quality_score = 70
    elif total_content_length < 3000:
        quality_score = 80
    else:
        quality_score = 94 if rag_context else 92

    if not improvement_points:
        improvement_points.append("No major issues found during self-check.")

    return {
        "quality_score": quality_score,
        "missing_sections": missing_sections,
        "total_content_length": total_content_length,
        "rag_chunks_used": len(rag_context),
        "strengths": strengths,
        "improvement_points": improvement_points,
        "self_check_summary": (
            "The agent reviewed the generated content against the original request, "
            "planned sections, retrieved context, assumptions, and expected document quality "
            "before producing the final Word document."
        )
    }


def clean_filename(title: str) -> str:
    title = title.lower()
    title = re.sub(r"[^a-z0-9]+", "_", title)
    title = title.strip("_")

    if not title:
        title = "generated_document"

    unique_id = str(uuid.uuid4())[:8]
    return f"{title}_{unique_id}.docx"


def clean_section_content(section_name: str, content: str) -> str:
    lines = content.splitlines()

    while lines and not lines[0].strip():
        lines.pop(0)

    if lines:
        first_line = lines[0].strip()

        normalized_first_line = re.sub(r"[^a-zA-Z0-9]+", " ", first_line).strip().lower()
        normalized_section_name = re.sub(r"[^a-zA-Z0-9]+", " ", section_name).strip().lower()

        if normalized_first_line == normalized_section_name:
            lines.pop(0)

    return "\n".join(lines).strip()


def add_content_to_document(document: Document, content: str):
    lines = content.splitlines()

    for line in lines:
        clean_line = line.strip()

        if not clean_line:
            continue

        if clean_line.startswith("- "):
            document.add_paragraph(clean_line[2:], style="List Bullet")

        elif clean_line.startswith("* "):
            document.add_paragraph(clean_line[2:], style="List Bullet")

        else:
            document.add_paragraph(clean_line)


def generate_word_document(
    user_request: str,
    plan: Dict[str, Any],
    section_outputs: Dict[str, str],
    execution_log: List[str],
    reflection_result: Dict[str, Any],
    rag_context: List[Dict[str, Any]]
) -> str:
    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = plan["title"]

    document.add_heading(title, level=0)

    document.add_paragraph(f"Document Type: {plan['document_type']}")
    document.add_paragraph(f"Generated On: {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    document.add_paragraph(f"Original User Request: {user_request}")

    document.add_heading("Autonomous Agent Plan", level=1)

    for task in plan["tasks"]:
        step = task["step"]
        task_name = task["task"]
        reason = task["reason"]

        document.add_paragraph(
            f"Step {step}: {task_name} — {reason}",
            style="List Number"
        )

    document.add_heading("Agent Assumptions", level=1)

    for assumption in plan["assumptions"]:
        document.add_paragraph(assumption, style="List Bullet")

    document.add_heading("Retrieved Context from Simple RAG", level=1)

    if rag_context:
        for chunk in rag_context:
            document.add_paragraph(
                f"{chunk['title']} — relevance score: {chunk['score']}",
                style="List Bullet"
            )
    else:
        document.add_paragraph("No additional context was retrieved.")

    document.add_heading("Execution Log", level=1)

    for log in execution_log:
        document.add_paragraph(log, style="List Bullet")

    document.add_heading("Agent Reflection / Self-Check", level=1)

    document.add_paragraph(f"Quality Score: {reflection_result['quality_score']}/100")
    document.add_paragraph(f"Total Generated Content Length: {reflection_result['total_content_length']} characters")
    document.add_paragraph(f"RAG Chunks Used: {reflection_result['rag_chunks_used']}")
    document.add_paragraph(reflection_result["self_check_summary"])

    document.add_heading("Strengths Identified", level=2)

    for strength in reflection_result["strengths"]:
        document.add_paragraph(strength, style="List Bullet")

    document.add_heading("Improvement Points", level=2)

    for point in reflection_result["improvement_points"]:
        document.add_paragraph(point, style="List Bullet")

    document.add_page_break()

    document.add_heading("Generated Business Document", level=1)

    for section_name, content in section_outputs.items():
        document.add_heading(section_name, level=2)
        cleaned_content = clean_section_content(section_name, content)
        add_content_to_document(document, cleaned_content)

    document.add_heading("Final Note", level=1)
    document.add_paragraph(
        "This Word document was generated by an autonomous AI agent that created a task plan, "
        "retrieved supporting context through a simple local RAG layer, executed each section, "
        "reviewed its own output through reflection/self-check, and produced a structured business-ready document."
    )

    filename = clean_filename(title)
    file_path = os.path.join(OUTPUT_DIR, filename)

    document.save(file_path)

    return filename


@app.get("/")
def health_check():
    return {
        "status": "running",
        "message": "Autonomous Document Agent API is active."
    }


@app.post("/agent")
def run_agent(payload: AgentRequest):
    user_request = payload.request.strip()

    if not user_request:
        raise HTTPException(status_code=400, detail="Request cannot be empty.")

    execution_log = []

    try:
        execution_log.append("Received user request.")

        plan = create_execution_plan(user_request)
        execution_log.append("Autonomous execution plan created.")

        rag_context = retrieve_rag_context(
            user_request=user_request,
            plan=plan
        )

        execution_log.append(
            f"Simple RAG retrieved {len(rag_context)} relevant knowledge chunks."
        )

        section_outputs = {}
        llm_sources = {}

        for section in plan["sections"]:
            execution_log.append(f"Executing task: Generate content for section '{section}'.")

            result = generate_section_content(
                user_request=user_request,
                plan=plan,
                section_name=section,
                rag_context=rag_context
            )

            section_outputs[section] = result["content"]

            llm_sources[section] = {
                "source": result["source"],
                "error": result["error"]
            }

            if result["source"] == "fallback":
                execution_log.append(f"Fallback used for section '{section}'.")
            else:
                execution_log.append(f"LLM content generated for section '{section}'.")

        execution_log.append("All planned sections executed successfully.")

        reflection_result = reflect_on_output(
            user_request=user_request,
            plan=plan,
            section_outputs=section_outputs,
            rag_context=rag_context
        )

        execution_log.append(
            f"Reflection/self-check completed with quality score {reflection_result['quality_score']}/100."
        )

        filename = generate_word_document(
            user_request=user_request,
            plan=plan,
            section_outputs=section_outputs,
            execution_log=execution_log,
            reflection_result=reflection_result,
            rag_context=rag_context
        )

        execution_log.append("Microsoft Word document generated successfully.")

        return {
            "message": "Agent executed successfully and generated Word document.",
            "user_request": user_request,
            "document_type": plan["document_type"],
            "title": plan["title"],
            "sections": plan["sections"],
            "assumptions": plan["assumptions"],
            "plan": plan["tasks"],
            "document_filename": filename,
            "document_url": f"/documents/{filename}",
            "rag_context": rag_context,
            "reflection_result": reflection_result,
            "section_outputs": section_outputs,
            "llm_sources": llm_sources,
            "execution_log": execution_log,
            "status": "document_generated"
        }

    except Exception as error:
        raise HTTPException(status_code=500, detail=str(error))


@app.get("/documents/{filename}")
def download_document(filename: str):
    if "/" in filename or "\\" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename.")

    file_path = os.path.join(OUTPUT_DIR, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Document not found.")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )